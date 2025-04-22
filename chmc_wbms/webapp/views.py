import base64, re
import os
import hashlib
import pythoncom
import json
import calendar
from win32com.client import Dispatch
from io import BytesIO
from PIL import Image
from django.core.files.base import ContentFile
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.hashers import make_password
from django.contrib.auth.decorators import user_passes_test, login_required
from django.views.decorators.http import require_GET, require_http_methods
from django.contrib.sessions.backends.db import SessionStore
from django.db.models import Sum, Count, Q, Prefetch, CharField
from webapp.models import Appointment, Payment, CustomUser, ServiceType, Patient, Examination, AppointmentTimeSlot, TimeSlot
from django.utils import timezone
from django.utils.crypto import get_random_string
from django.http import HttpResponseRedirect, JsonResponse, FileResponse, HttpResponse
from django.urls import reverse
from datetime import datetime, timedelta, time
from .forms import UserCreationForm, EditAccountForm, EditProfileForm, AppointmentForm, ExaminationForm, UploadEditedDocumentForm, UploadResultImageForm, EditExaminationForm, PatientEditForm
from django.views.decorators.csrf import csrf_protect, csrf_exempt
from docx import Document
from docx.shared import Inches
from django.conf import  settings
from django.core.exceptions import ObjectDoesNotExist
from django.core.paginator import Paginator
from django.core.files.storage import default_storage
from urllib.parse import urlencode
from django.db.models.functions import Cast
from decimal import Decimal

@csrf_protect
def admin_login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)

        if user is not None and user.is_superuser:
            login(request, user)
            # Set custom session data (Django manages session storage automatically)
            session_key = f"admin_session_{get_random_string(32)}"
            request.session['session_key'] = session_key
            request.session['user_id'] = user.id
            response = HttpResponseRedirect(reverse('admin_dashboard'))
            response.set_cookie('admin_session', session_key, httponly=True)
            return response
        else:
            messages.error(request, "Invalid credentials or unauthorized access.")
    return render(request, 'admin/admin_login.html')


def user_login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)

            # Check if user is an employee or a clinic doctor
            if user is not None:
                if (hasattr(user, 'is_employee') and user.is_employee) or \
                   (hasattr(user, 'is_clinic_doctor') and user.is_clinic_doctor):
                    login(request, user)
                    # Redirect to employee dashboard
                    return HttpResponseRedirect(reverse('employee_dashboard'))
                elif hasattr(user, 'is_associated_doctor') and user.is_associated_doctor:
                    login(request, user)
                    # Redirect to associated doctor dashboard
                    return HttpResponseRedirect(reverse('assocdoc_dashboard'))
            else:
                messages.error(request, "Invalid credentials or unauthorized access.")
    else:
        form = AuthenticationForm()

    return render(request, 'employee/employee_login.html', {'form': form})

@user_passes_test(lambda u: u.is_superuser)
def manage_account_view(request):
    # Filter out the superuser accounts and get only employees and associated doctors
    accounts = CustomUser.objects.exclude(is_superuser=True).filter(Q(is_employee=True) | Q(is_associated_doctor=True) | Q(is_clinic_doctor=True)
    )

    # Context to pass to the template
    context = {
        'accounts': accounts,
    }
    return render(request, 'admin/manage_accounts.html', context)

@user_passes_test(lambda u: u.is_superuser)
def patients_list_view(request, patient_id=None):
    # Get all patients but only fetch the current page's patients
    patients = Patient.objects.order_by('id')
    paginator = Paginator(patients, 12)  # 12 patients per page
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    search_query = request.GET.get('search', '').strip()

    selected_patient = None
    examinations = None

    if search_query:
        patients = patients.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(middle_name__icontains=search_query) |
            Q(id__icontains=search_query.replace("PID-", ""))
        )

    if patient_id:
        selected_patient = get_object_or_404(Patient, id=patient_id)
        # Fetch only examinations for the selected patient
        examinations = Examination.objects.filter(patient=selected_patient) \
            .select_related('attending_doctor') \
            .prefetch_related('service_types', 'payment_set') \
            .order_by('-date_created')

    context = {
        'account': request.user,
        'service_types': ServiceType.objects.all(),
        'selected_patient': selected_patient,
        'examinations': examinations,
        'page_obj': page_obj,  # Use this in the template
        'search_query' : search_query,
    }
    
    return render(request, 'admin/patients_list.html', context)


@user_passes_test(lambda u: u.is_superuser)
def admin_dashboard_view(request):
    account = request.user
    today = timezone.now().date()
    
    # Calculate dashboard metrics using Payment model
    # Daily income
    daily_income = Payment.objects.filter(
        date__date=today,
        status='Paid'
    ).aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    
    # Weekly income (last 7 days including today)
    start_of_week = today - timedelta(days=6)
    weekly_income = Payment.objects.filter(
        date__date__gte=start_of_week,
        status='Paid'
    ).aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    
    # Monthly income
    start_of_month = today.replace(day=1)
    monthly_income = Payment.objects.filter(
        date__date__gte=start_of_month,
        status='Paid'
    ).aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    
    # Count daily patients served (distinct patients)
    daily_patients = Examination.objects.filter(
        date_created__date=today
    ).values('patient').distinct().count()
    
    # Count daily examinations
    daily_examinations = Examination.objects.filter(
        date_created__date=today
    ).count()

    # Handle date filtering
    selected_date = request.GET.get('date')
    try:
        selected_date = datetime.strptime(selected_date, '%Y-%m-%d').date() if selected_date else None
    except (ValueError, TypeError):
        selected_date = None

    # Base queryset - will be filtered if date is specified
    appointments = Appointment.objects.all().order_by('appointment_date')
    
    # Apply date filter if specific date is selected (not None and not empty string)
    if selected_date:
        appointments = appointments.filter(appointment_date=selected_date)

    paginator = Paginator(appointments, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    base_query = ''
    if selected_date:
        base_query = f'date={selected_date.strftime("%Y-%m-%d")}'

    service_types = ServiceType.objects.all()

    if request.method == 'POST':
        # Handle appointment creation
        if 'add_appointment' in request.POST:
            form = AppointmentForm(request.POST)
            if form.is_valid():
                appointment = form.save(commit=False)
                appointment.save()
                form.save_m2m()  # Save service types
                
                # Handle time slots
                time_slot_ids = request.POST.getlist('appointment_time')
                for slot_id in time_slot_ids:
                    time_slot = TimeSlot.objects.get(id=slot_id)
                    AppointmentTimeSlot.objects.create(
                        appointment=appointment,
                        time_slot=time_slot,
                        date=appointment.appointment_date
                    )
                
                # Handle special slot if selected
                special_slot_type = request.POST.get('special_slot')
                if special_slot_type:
                    special_slot = TimeSlot.objects.get(
                        is_special=True,
                        special_type=special_slot_type
                    )
                    AppointmentTimeSlot.objects.create(
                        appointment=appointment,
                        time_slot=special_slot,
                        date=appointment.appointment_date
                    )
                
                messages.success(request, 'Appointment scheduled successfully.')
                return redirect('admin_dashboard')
        
        # Handle appointment update
        elif 'update_appointment' in request.POST:
            appointment_id = request.POST.get('appointment_id')
            try:
                appointment = Appointment.objects.get(id=appointment_id)
                form = AppointmentForm(request.POST, instance=appointment)
                if form.is_valid():
                    updated_appointment = form.save(commit=False)
                    updated_appointment.save()
                    form.save_m2m()  # Update service types
                    
                    # Clear existing time slots
                    appointment.appointmenttimeslot_set.all().delete()
                    
                    # Add new regular time slots
                    time_slot_ids = request.POST.getlist('appointment_time')
                    for slot_id in time_slot_ids:
                        time_slot = TimeSlot.objects.get(id=slot_id)
                        AppointmentTimeSlot.objects.create(
                            appointment=appointment,
                            time_slot=time_slot,
                            date=appointment.appointment_date
                        )
                    
                    # Handle special slot if selected
                    special_slot_type = request.POST.get('special_slot')
                    if special_slot_type:
                        special_slot = TimeSlot.objects.get(
                            is_special=True,
                            special_type=special_slot_type
                        )
                        AppointmentTimeSlot.objects.create(
                            appointment=appointment,
                            time_slot=special_slot,
                            date=appointment.appointment_date
                        )
                    
                    messages.success(request, 'Appointment updated successfully.')
                    return redirect('admin_dashboard')
            
            except Exception as e:
                messages.error(request, f'Error updating appointment: {str(e)}')
        
        # Handle appointment deletion
        elif 'delete_appointment' in request.POST:
            appointment_id = request.POST.get('appointment_id')
            try:
                appointment = Appointment.objects.get(id=appointment_id)
                appointment.delete()
                messages.success(request, 'Appointment deleted successfully.')
                return redirect('admin_dashboard')
            except Appointment.DoesNotExist:
                messages.error(request, 'Appointment not found')
            except Exception as e:
                messages.error(request, f'Error deleting appointment: {str(e)}')

    else:
        form = AppointmentForm()

    context = {
        'daily_income': daily_income,
        'weekly_income': weekly_income,
        'monthly_income': monthly_income,
        'daily_patients': daily_patients,
        'daily_examinations': daily_examinations,
        'appointments': appointments,
        'service_types': service_types,
        'form': form,
        'account': account,
        'selected_date': selected_date.strftime('%Y-%m-%d') if selected_date else '',
        'page_obj': page_obj,
        'base_query': base_query,
    }
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return render(request, 'admin/partials/appointments_table.html', context)

    return render(request, 'admin/admin_dashboard.html', context)

@user_passes_test(lambda u: u.is_superuser)
def admin_examination_view(request):
# Fetch examinations along with related patient and payment details
    account = request.user
    examinations = (
        Examination.objects.select_related('patient', 'attending_doctor')
        .prefetch_related('service_types', Prefetch('payment_set'))
        .order_by('-date_created')  # Order by latest examination first
    )
     
    search_query = request.GET.get('search', '').strip()
    date_filter = request.GET.get('date')
    
    # Apply filters first
    if search_query:
        examinations = examinations.filter(
            Q(patient__first_name__icontains=search_query) |
            Q(patient__last_name__icontains=search_query) |
            Q(patient__middle_name__icontains=search_query) |
            Q(patient__id__icontains=search_query)  # Make sure this field is searchable
        )
    
    if date_filter:
        try:
            filter_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
            examinations = examinations.filter(date_created__date=filter_date)
        except ValueError:
            pass
    service_types = ServiceType.objects.all()

    paginator = Paginator(examinations, 10)  # or however many per page
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    
    context = {
        'account' : account,
        'examinations': page_obj,
        'service_types': service_types,
        'page_obj' : page_obj,
        'search_query' : search_query,
        'date_filter' : date_filter,
    }
    return render(request, 'admin/examination.html', context)

@user_passes_test(lambda u: u.is_superuser)
def admin_document_results_view(request):
    account = request.user
    search_query = request.GET.get('search', '').strip()
    
    examinations = Examination.objects.all().order_by('-date_created')
    years = Examination.objects.filter(edited_document__isnull=False).dates('date_created', 'year', order='DESC')

    selected_year = request.GET.get('year')
    selected_month = request.GET.get('month')
    selected_day = request.GET.get('day')
    selected_filter = request.GET.get('filter', 'all')

    documents = Examination.objects.filter(edited_document__isnull=False)
    months = {}
    days = {}
    page_obj = None

    # Search functionality (filters by filename, patient, or doctor)
    if search_query:
        documents = [
            doc for doc in Examination.objects.all() 
            if doc.edited_document and search_query.lower() in doc.filename.lower()
        ]
        # If it's an AJAX/HTMX request, return only search results
        if request.headers.get('HX-Request') == 'true':
            paginator = Paginator(documents, 10)
            page_number = request.GET.get("page", 1)
            page_obj = paginator.get_page(page_number)
            return render(request, 'admin/partials/search_results.html', {
                'documents': page_obj,
                'search_query': search_query,
            })

    # Year/Month/Day filtering
    if selected_year:
        year_int = int(selected_year)
        documents = documents.filter(date_created__year=year_int)
        
        # Only show months that have documents
        months = {
            calendar.month_name[i]: documents.filter(date_created__month=i).exists()
            for i in range(1, 13) if documents.filter(date_created__month=i).exists()
        }

        if selected_month:
            month_int = next(i for i, m in enumerate(calendar.month_name) if m == selected_month)
            documents = documents.filter(date_created__month=month_int)
            
            # Only show days that have documents
            days_in_month = documents.dates('date_created', 'day')
            days = {day.day: True for day in days_in_month}

            if selected_day:
                documents = documents.filter(date_created__day=int(selected_day))

    # Pagination
    paginator = Paginator(documents, 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)

    # Base query for pagination links (excludes 'page' param)
    querydict = request.GET.copy()
    if 'page' in querydict:
        querydict.pop('page')
    base_query = urlencode(querydict)

    context = {
        'years': [y.year for y in years],
        'selected_year': selected_year,
        'selected_month': selected_month,
        'selected_day': selected_day,
        'selected_filter': selected_filter,
        'months': months,
        'days': days,
        'documents': documents,
        'page_obj': page_obj,
        'account': account,
        'base_query': base_query,
        'search_query': search_query,
        'examinations' : examinations,
    }
    return render(request, 'admin/document_results.html', context)

@login_required
def admin_add_examination(request):
    account = request.user

    if request.method == 'POST':
        # Handle patient search (AJAX request)
        if 'search_patient' in request.POST:
            search_query = request.POST.get('search_patient', '').strip()
            # Filter patients by name or ID
            patients = Patient.objects.filter(
                Q(first_name__istartswith=search_query) |
                Q(last_name__istartswith=search_query) |
                Q(middle_name__istartswith=search_query) |
                Q(id__icontains=search_query)
            )
            # Prepare patient data for JSON response
            patient_list = [
                {
                    'id': patient.id,
                    'first_name': patient.first_name,
                    'last_name': patient.last_name,
                    'middle_name': patient.middle_name or '',
                    'age': patient.age,
                    'sex': patient.sex,
                    'address': patient.address,
                    'contact_number': patient.contact_number,
                    'image_url': patient.image.url if patient.image else None,
                }
                for patient in patients
            ]
            return JsonResponse({'patients': patient_list})

        # Handle form submission for adding an examination
        form = ExaminationForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                # Check if a patient is selected from the search results
                patient_id = request.POST.get('patient_id')
                if patient_id:
                    patient = get_object_or_404(Patient, id=patient_id)
                    
                    # If the patient has an existing image, remove it before adding a new one
                    if patient.image:
                        patient.image.delete()
                    
                    # Handle the new image from the webcam
                    image_data = request.POST.get('image')  # Captured image data from the form
                    if image_data:
                        # Convert base64 image data to a Django ImageFile
                        format, imgstr = image_data.split(';base64,')  # Get base64 string
                        ext = format.split('/')[1]  # Extract file extension (png, jpeg, etc.)
                        image_data = ContentFile(base64.b64decode(imgstr), name=f"patient_{patient.id}_image.{ext}")
                        
                        # Save the new image to the patient instance
                        patient.image = image_data
                        patient.save()
                else:
                    # Create a new patient if not selected
                    first_name = form.cleaned_data['first_name']
                    last_name = form.cleaned_data['last_name']
                    middle_name = form.cleaned_data['middle_name']
                    age = form.cleaned_data['age']
                    sex = form.cleaned_data['sex']
                    address = form.cleaned_data['address']
                    contact_number = form.cleaned_data['contact_number']

                    patient = Patient.objects.create(
                        first_name=first_name,
                        last_name=last_name,
                        middle_name=middle_name,
                        age=age,
                        sex=sex,
                        address=address,
                        contact_number=contact_number,
                    )

                # Create the examination
                examination = Examination.objects.create(
                    patient=patient,
                    attending_doctor=form.cleaned_data['attending_doctor']
                )
                examination.service_types.set(form.cleaned_data['service_types'])

                # Create payment record
                Payment.objects.create(
                    examination=examination,
                    method=form.cleaned_data['method'],
                    amount=form.cleaned_data['amount'],
                    status=form.cleaned_data['status']
                )

                # Generate the document
                generate_examination_document(examination)

                return redirect('admin_examination')  # Redirect to success page
            except Exception as e:
                return render(request, 'admin/add_examination.html', {
                    'form': form,
                    'account': account,
                    'error': f"An error occurred: {e}"
                })
    else:
        form = ExaminationForm()

    return render(request, 'admin/add_examination.html', {
        'form': form,
        'account': account,
    })

@login_required
def admin_edit_patient(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)
    if request.method == 'POST':
        form = PatientEditForm(request.POST, request.FILES, instance=patient)
        if form.is_valid():
            form.save()
            return redirect('admin_patients_list_with_id', patient_id=patient.id)  # Redirect after saving changes
    else:
        form = PatientEditForm(instance=patient)
    
    account = request.user
    context = {
        'form': form,
        'patient': patient,
        'account': account,
    }
    return render(request, 'admin/edit_patient.html', context)

@user_passes_test(lambda u: u.is_superuser)
def delete_examination_view(request, pk):
    exam = get_object_or_404(Examination, pk=pk)
    exam.delete()
    return redirect('admin_examination')

@user_passes_test(lambda u: u.is_superuser)
def delete_patient_view(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    patient.delete()
    return redirect('admin_patients_list')

@user_passes_test(lambda u: u.is_superuser)
def delete_edited_document_view(request, pk):
    exam = get_object_or_404(Examination, pk=pk)
    if exam.edited_document:
        exam.edited_document.delete(save=False)  # Delete file from storage
        exam.edited_document = None
        exam.save()
    return redirect(request.META.get('HTTP_REFERER', 'admin_document_results'))

@user_passes_test(lambda u: u.is_superuser)
def admin_logout_view(request):
    logout(request)  # This logs out the user
    return redirect('admin_login') 

@user_passes_test(lambda u: u.is_employee or u.is_clinic_doctor)
def employee_logout_view(request):
    logout(request)  # This logs out the user
    return redirect('user_login') 

@user_passes_test(lambda u: u.is_associated_doctor)
def assocdoc_logout_view(request):
    logout(request)  # This logs out the user
    return redirect('user_login') 

@user_passes_test(lambda u: u.is_superuser)
def create_account_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST, request.FILES)  # Include `request.FILES` for file uploads
        if form.is_valid():
            # Save the new user
            user = form.save(commit=False)
            user.username = form.cleaned_data['email']
            user.password = make_password(form.cleaned_data['password'])
            user.first_name = form.cleaned_data['first_name']
            user.last_name = form.cleaned_data['last_name']
            user.email = form.cleaned_data['email']

            # Determine account type
            account_type = request.POST.get('account_type')  # Get the selected account type
            if account_type == 'employee':
                user.is_employee = True
            elif account_type == 'assoc_doctor':
                user.is_associated_doctor = True
                user.signature_image = form.cleaned_data.get('signature_image')  # Add signature image
            elif account_type == 'clinic_doctor':
                # Check if a clinic doctor already exists
                if CustomUser.objects.filter(is_clinic_doctor=True).exists():
                    messages.error(request, "A clinic doctor account already exists.")
                    return redirect('create_account')  # Redirect back to the form
                user.is_clinic_doctor = True
                user.signature_image = form.cleaned_data.get('signature_image')  # Add signature image

            # Save the user to the database
            user.save()

            messages.success(request, "Account created successfully!")
            return redirect('admin_dashboard')
        else:
            messages.error(request, "Error creating account. Please check the details.")

    else:
        form = UserCreationForm()
    return render(request, 'admin/admin_dashboard.html', {'form': form})

@user_passes_test(lambda u: u.is_employee or u.is_clinic_doctor)
def employee_dashboard_view(request):
    account = request.user
    today = timezone.now().date()
    
    # Calculate dashboard metrics using Payment model
    # Daily income
    daily_income = Payment.objects.filter(
        date__date=today,
        status='Paid'
    ).aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    
    # Weekly income (last 7 days including today)
    start_of_week = today - timedelta(days=6)
    weekly_income = Payment.objects.filter(
        date__date__gte=start_of_week,
        status='Paid'
    ).aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    
    # Monthly income
    start_of_month = today.replace(day=1)
    monthly_income = Payment.objects.filter(
        date__date__gte=start_of_month,
        status='Paid'
    ).aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    
    # Count daily patients served (distinct patients)
    daily_patients = Examination.objects.filter(
        date_created__date=today
    ).values('patient').distinct().count()
    
    # Count daily examinations
    daily_examinations = Examination.objects.filter(
        date_created__date=today
    ).count()

    # Handle date filtering
    selected_date = request.GET.get('date')
    try:
        selected_date = datetime.strptime(selected_date, '%Y-%m-%d').date() if selected_date else None
    except (ValueError, TypeError):
        selected_date = None

    # Base queryset - will be filtered if date is specified
    appointments = Appointment.objects.all().order_by('appointment_date')
    
    # Apply date filter if specific date is selected (not None and not empty string)
    if selected_date:
        appointments = appointments.filter(appointment_date=selected_date)

    paginator = Paginator(appointments, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    base_query = ''
    if selected_date:
        base_query = f'date={selected_date.strftime("%Y-%m-%d")}'

    service_types = ServiceType.objects.all()

    if request.method == 'POST':
        # Handle appointment creation
        if 'add_appointment' in request.POST:
            form = AppointmentForm(request.POST)
            if form.is_valid():
                appointment = form.save(commit=False)
                appointment.save()
                form.save_m2m()  # Save service types
                
                # Handle time slots
                time_slot_ids = request.POST.getlist('appointment_time')
                for slot_id in time_slot_ids:
                    time_slot = TimeSlot.objects.get(id=slot_id)
                    AppointmentTimeSlot.objects.create(
                        appointment=appointment,
                        time_slot=time_slot,
                        date=appointment.appointment_date
                    )
                
                # Handle special slot if selected
                special_slot_type = request.POST.get('special_slot')
                if special_slot_type:
                    special_slot = TimeSlot.objects.get(
                        is_special=True,
                        special_type=special_slot_type
                    )
                    AppointmentTimeSlot.objects.create(
                        appointment=appointment,
                        time_slot=special_slot,
                        date=appointment.appointment_date
                    )
                
                messages.success(request, 'Appointment scheduled successfully.')
                return redirect('employee_dashboard')
        
        # Handle appointment update
        elif 'update_appointment' in request.POST:
            appointment_id = request.POST.get('appointment_id')
            try:
                appointment = Appointment.objects.get(id=appointment_id)
                form = AppointmentForm(request.POST, instance=appointment)
                if form.is_valid():
                    updated_appointment = form.save(commit=False)
                    updated_appointment.save()
                    form.save_m2m()  # Update service types
                    
                    # Clear existing time slots
                    appointment.appointmenttimeslot_set.all().delete()
                    
                    # Add new regular time slots
                    time_slot_ids = request.POST.getlist('appointment_time')
                    for slot_id in time_slot_ids:
                        time_slot = TimeSlot.objects.get(id=slot_id)
                        AppointmentTimeSlot.objects.create(
                            appointment=appointment,
                            time_slot=time_slot,
                            date=appointment.appointment_date
                        )
                    
                    # Handle special slot if selected
                    special_slot_type = request.POST.get('special_slot')
                    if special_slot_type:
                        special_slot = TimeSlot.objects.get(
                            is_special=True,
                            special_type=special_slot_type
                        )
                        AppointmentTimeSlot.objects.create(
                            appointment=appointment,
                            time_slot=special_slot,
                            date=appointment.appointment_date
                        )
                    
                    messages.success(request, 'Appointment updated successfully.')
                    return redirect('employee_dashboard')
            
            except Exception as e:
                messages.error(request, f'Error updating appointment: {str(e)}')
        
        # Handle appointment deletion
        elif 'delete_appointment' in request.POST:
            appointment_id = request.POST.get('appointment_id')
            try:
                appointment = Appointment.objects.get(id=appointment_id)
                appointment.delete()
                messages.success(request, 'Appointment deleted successfully.')
                return redirect('employee_dashboard')
            except Appointment.DoesNotExist:
                messages.error(request, 'Appointment not found')
            except Exception as e:
                messages.error(request, f'Error deleting appointment: {str(e)}')

    else:
        form = AppointmentForm()

    context = {
        'daily_income': daily_income,
        'weekly_income': weekly_income,
        'monthly_income': monthly_income,
        'daily_patients': daily_patients,
        'daily_examinations': daily_examinations,
        'appointments': appointments,
        'service_types': service_types,
        'form': form,
        'account': account,
        'selected_date': selected_date.strftime('%Y-%m-%d') if selected_date else '',
        'page_obj': page_obj,
        'base_query': base_query,
    }
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return render(request, 'employee/partials/appointments_table.html', context)

    return render(request, 'employee/employee_dashboard.html', context)

@require_GET
def get_available_time_slots(request):
    date_str = request.GET.get('date')
    if not date_str:
        return JsonResponse({'error': 'Date required'}, status=400)
    
    try:
        date = datetime.strptime(date_str, '%Y-%m-%d').date()
    except ValueError:
        return JsonResponse({'error': 'Invalid date format'}, status=400)
    
    # Get ALL booked slot IDs for this date
    booked_slot_ids = AppointmentTimeSlot.objects.filter(
        date=date
    ).values_list('time_slot_id', flat=True)
    
    # Get special slot bookings
    booked_special = AppointmentTimeSlot.objects.filter(
        date=date,
        time_slot__is_special=True
    ).select_related('time_slot')
    
    # Get ALL active slots
    all_slots = TimeSlot.objects.filter(is_active=True)
    
    # Filter available slots
    available_slots = []
    for slot in all_slots:
        # Skip if slot is booked
        if slot.id in booked_slot_ids:
            continue
            
        # Check special slot conflicts
        if not slot.is_special:
            slot_available = True
            for special_booking in booked_special:
                if (special_booking.time_slot.special_type == 'MORNING' and 
                    slot.start_time >= time(8, 0) and slot.start_time < time(12, 0)):
                    slot_available = False
                    break
                elif (special_booking.time_slot.special_type == 'AFTERNOON' and 
                      slot.start_time >= time(14, 0) and slot.start_time < time(17, 0)):
                    slot_available = False
                    break
                elif special_booking.time_slot.special_type == 'DAY':
                    slot_available = False
                    break
            
            if not slot_available:
                continue
                
        available_slots.append(slot)
    
    # Format response
    slots_data = [{
        'id': slot.id,
        'start_time': slot.start_time.strftime('%I:%M %p'),
        'end_time': slot.end_time.strftime('%I:%M %p'),
        'is_special': slot.is_special,
        'special_type': slot.special_type if slot.is_special else None
    } for slot in available_slots]
    
    return JsonResponse({
        'available_slots': slots_data,
        'date': date_str
    })

def get_appointment_details(request):
    if request.method == 'GET' and request.GET.get('appointment_id'):
        try:
            appointment = Appointment.objects.get(id=request.GET['appointment_id'])
            
            # Get current service types
            service_types = list(appointment.service_types.values_list('id', flat=True))
            
            # Get all time slots for this appointment
            appointment_slots = appointment.appointmenttimeslot_set.select_related('time_slot').all()
            
            # Separate regular and special slots
            current_time_slots = []
            current_special_slots = []
            
            for slot in appointment_slots:
                if slot.time_slot.is_special:
                    current_special_slots.append(slot.time_slot.special_type)
                else:
                    current_time_slots.append(slot.time_slot.id)
            
            # Get available time slots for the appointment date
            available_time_slots = []
            for slot in TimeSlot.objects.filter(is_active=True, is_special=False):
                start_time = slot.start_time.strftime('%I:%M %p').lstrip('0')
                end_time = slot.end_time.strftime('%I:%M %p').lstrip('0')
                available_time_slots.append({
                    'id': slot.id,
                    'start_time': start_time,
                    'end_time': end_time,
                    'display': f"{start_time} - {end_time}"
                })
            
            # Get available special slots
            available_special_slots = []
            special_slot_types = [
                ('MORNING', 'Morning (8AM-12PM)'),
                ('AFTERNOON', 'Afternoon (2PM-5PM)'),
                ('DAY', 'Whole Day (8AM-5PM)')
            ]
            
            for slot_type, display_name in special_slot_types:
                # Check if this special slot type is already booked by another appointment
                is_available = not AppointmentTimeSlot.objects.filter(
                    date=appointment.appointment_date,
                    time_slot__special_type=slot_type
                ).exclude(appointment=appointment).exists()
                
                if is_available:
                    available_special_slots.append({
                        'special_type': slot_type,
                        'display_name': display_name
                    })
            
            return JsonResponse({
                'success': True,
                'appointment': {
                    'client_name': appointment.client_name,
                    'appointment_date': appointment.appointment_date.strftime('%Y-%m-%d'),
                    'description': appointment.description,
                    'service_types': service_types,
                },
                'current_time_slots': current_time_slots,
                'current_special_slots': current_special_slots,
                'available_time_slots': list(available_time_slots),
                'available_special_slots': available_special_slots,
            })
        except Appointment.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Appointment not found'})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    return JsonResponse({'success': False, 'message': 'Invalid request'})

@user_passes_test(lambda u: u.is_superuser)
def edit_account_view(request, account_id):
    account = get_object_or_404(CustomUser, id=account_id)
    form = EditAccountForm(request.POST or None, request.FILES or None, instance=account)
    
    if request.method == 'POST':
        if form.is_valid():
            form.save()
            messages.success(request, 'Account updated successfully.')
            return redirect('manage_accounts')
        else:
            messages.error(request, 'Please correct the error below.')
    
    return render(request, 'admin/edit_account.html', {'form': form, 'account': account})

@user_passes_test(lambda u: u.is_superuser)
def delete_account_view(request, account_id):
    account = get_object_or_404(CustomUser, id=account_id)
    if request.method == 'POST':
        account.delete()
        messages.success(request, 'Account deleted successfully.')
        return redirect('manage_accounts')
    
    return render(request, 'admin/delete_account.html', {'account': account})

@user_passes_test(lambda u: u.is_employee or u.is_clinic_doctor)
def edit_profile_view(request, account_id):
    account = get_object_or_404(CustomUser, id=account_id)
    form = EditProfileForm(request.POST or None, request.FILES or None, instance=account)
    
    if request.method == 'POST':
        if form.is_valid():
            form.save()
            messages.success(request, 'Account updated successfully.')
            return redirect('employee_dashboard')
        else:
            messages.error(request, 'Please correct the error below.')
    
    return render(request, 'employee/edit_profile.html', {'form': form, 'account': account})

def search_patients_list(request):
    query = request.GET.get('q', '').strip()
    if not query:
        return JsonResponse({'patients': []})

    patients = Patient.objects.filter(
        Q(first_name__icontains=query) |
        Q(last_name__icontains=query) |
        Q(middle_name__icontains=query) |
        Q(id__icontains=query.replace("PID-", ""))
    ).order_by('last_name', 'first_name')[:10]

    results = [{
        'id': patient.id,
        'text': f"{patient.patient_full_name_last_name_start()} ({patient.get_formatted_id()})",
        'formatted_id': patient.get_formatted_id()
    } for patient in patients]

    return JsonResponse({'patients': results})

@user_passes_test(lambda u: u.is_employee or u.is_clinic_doctor)
def employee_patients_list_view(request, patient_id=None):
    # Get all patients but only fetch the current page's patients
    patients = Patient.objects.order_by('id')
    paginator = Paginator(patients, 12)  # 12 patients per page
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    search_query = request.GET.get('search', '').strip()

    selected_patient = None
    examinations = None

    if search_query:
        patients = patients.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(middle_name__icontains=search_query) |
            Q(id__icontains=search_query.replace("PID-", ""))
        )

    if patient_id:
        selected_patient = get_object_or_404(Patient, id=patient_id)
        # Fetch only examinations for the selected patient
        examinations = Examination.objects.filter(patient=selected_patient) \
            .select_related('attending_doctor') \
            .prefetch_related('service_types', 'payment_set') \
            .order_by('-date_created')

    context = {
        'account': request.user,
        'service_types': ServiceType.objects.all(),
        'selected_patient': selected_patient,
        'examinations': examinations,
        'page_obj': page_obj,  # Use this in the template
        'search_query' : search_query,
    }
    
    return render(request, 'employee/patients_list.html', context)

@user_passes_test(lambda u: u.is_employee or u.is_clinic_doctor)
def document_results_view(request):
    account = request.user
    search_query = request.GET.get('search', '').strip()
    
    examinations = Examination.objects.all().order_by('-date_created')
    years = Examination.objects.filter(edited_document__isnull=False).dates('date_created', 'year', order='DESC')

    selected_year = request.GET.get('year')
    selected_month = request.GET.get('month')
    selected_day = request.GET.get('day')
    selected_filter = request.GET.get('filter', 'all')

    documents = Examination.objects.filter(edited_document__isnull=False)
    months = {}
    days = {}
    page_obj = None

    # Search functionality (filters by filename, patient, or doctor)
    if search_query:
        documents = [
            doc for doc in Examination.objects.all() 
            if doc.edited_document and search_query.lower() in doc.filename.lower()
        ]
        # If it's an AJAX/HTMX request, return only search results
        if request.headers.get('HX-Request') == 'true':
            paginator = Paginator(documents, 10)
            page_number = request.GET.get("page", 1)
            page_obj = paginator.get_page(page_number)
            return render(request, 'employee/partials/search_results.html', {
                'documents': page_obj,
                'search_query': search_query,
            })

    # Year/Month/Day filtering
    if selected_year:
        year_int = int(selected_year)
        documents = documents.filter(date_created__year=year_int)
        
        # Only show months that have documents
        months = {
            calendar.month_name[i]: documents.filter(date_created__month=i).exists()
            for i in range(1, 13) if documents.filter(date_created__month=i).exists()
        }

        if selected_month:
            month_int = next(i for i, m in enumerate(calendar.month_name) if m == selected_month)
            documents = documents.filter(date_created__month=month_int)
            
            # Only show days that have documents
            days_in_month = documents.dates('date_created', 'day')
            days = {day.day: True for day in days_in_month}

            if selected_day:
                documents = documents.filter(date_created__day=int(selected_day))

    # Pagination
    paginator = Paginator(documents, 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)

    # Base query for pagination links (excludes 'page' param)
    querydict = request.GET.copy()
    if 'page' in querydict:
        querydict.pop('page')
    base_query = urlencode(querydict)

    context = {
        'years': [y.year for y in years],
        'selected_year': selected_year,
        'selected_month': selected_month,
        'selected_day': selected_day,
        'selected_filter': selected_filter,
        'months': months,
        'days': days,
        'documents': documents,
        'page_obj': page_obj,
        'account': account,
        'base_query': base_query,
        'search_query': search_query,
        'examinations' : examinations,
    }
    return render(request, 'employee/document_results.html', context)

@user_passes_test(lambda u: u.is_employee or u.is_clinic_doctor)
def employee_examination_view(request):
    # Fetch examinations along with related patient and payment details
    account = request.user
    examinations = (
        Examination.objects.select_related('patient', 'attending_doctor')
        .prefetch_related('service_types', Prefetch('payment_set'))
        .order_by('-date_created')  # Order by latest examination first
    )
     
    search_query = request.GET.get('search', '').strip()
    date_filter = request.GET.get('date')
    
    # Apply filters first
    if search_query:
        examinations = examinations.filter(
            Q(patient__first_name__icontains=search_query) |
            Q(patient__last_name__icontains=search_query) |
            Q(patient__middle_name__icontains=search_query) |
            Q(patient__id__icontains=search_query)  # Make sure this field is searchable
        )
    
    if date_filter:
        try:
            filter_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
            examinations = examinations.filter(date_created__date=filter_date)
        except ValueError:
            pass
    service_types = ServiceType.objects.all()

    paginator = Paginator(examinations, 10)  # or however many per page
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    
    context = {
        'account' : account,
        'examinations': page_obj,
        'service_types': service_types,
        'page_obj' : page_obj,
        'search_query' : search_query,
        'date_filter' : date_filter,
    }
    return render(request, 'employee/examination.html', context)

@login_required
def add_examination(request):
    account = request.user

    if request.method == 'POST':
        # Handle patient search (AJAX request)
        if 'search_patient' in request.POST:
            search_query = request.POST.get('search_patient', '').strip()
            # Filter patients by name or ID
            patients = Patient.objects.filter(
                Q(first_name__istartswith=search_query) |
                Q(last_name__istartswith=search_query) |
                Q(middle_name__istartswith=search_query) |
                Q(id__icontains=search_query)
            )
            # Prepare patient data for JSON response
            patient_list = [
                {
                    'id': patient.id,
                    'first_name': patient.first_name,
                    'last_name': patient.last_name,
                    'middle_name': patient.middle_name or '',
                    'age': patient.age,
                    'sex': patient.sex,
                    'address': patient.address,
                    'contact_number': patient.contact_number,
                    'image_url': patient.image.url if patient.image else None,
                }
                for patient in patients
            ]
            return JsonResponse({'patients': patient_list})

        # Handle form submission for adding an examination
        form = ExaminationForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                # Check if a patient is selected from the search results
                patient_id = request.POST.get('patient_id')
                if patient_id:
                    patient = get_object_or_404(Patient, id=patient_id)
                    
                    # If the patient has an existing image, remove it before adding a new one
                    if patient.image:
                        patient.image.delete()
                    
                    # Handle the new image from the webcam
                    image_data = request.POST.get('image')  # Captured image data from the form
                    if image_data:
                        # Convert base64 image data to a Django ImageFile
                        format, imgstr = image_data.split(';base64,')  # Get base64 string
                        ext = format.split('/')[1]  # Extract file extension (png, jpeg, etc.)
                        image_data = ContentFile(base64.b64decode(imgstr), name=f"patient_{patient.id}_image.{ext}")
                        
                        # Save the new image to the patient instance
                        patient.image = image_data
                        patient.save()
                else:
                    # Create a new patient if not selected
                    first_name = form.cleaned_data['first_name']
                    last_name = form.cleaned_data['last_name']
                    middle_name = form.cleaned_data['middle_name']
                    age = form.cleaned_data['age']
                    sex = form.cleaned_data['sex']
                    address = form.cleaned_data['address']
                    contact_number = form.cleaned_data['contact_number']

                    patient = Patient.objects.create(
                        first_name=first_name,
                        last_name=last_name,
                        middle_name=middle_name,
                        age=age,
                        sex=sex,
                        address=address,
                        contact_number=contact_number,
                    )

                # Create the examination
                examination = Examination.objects.create(
                    patient=patient,
                    attending_doctor=form.cleaned_data['attending_doctor']
                )
                examination.service_types.set(form.cleaned_data['service_types'])

                # Create payment record
                Payment.objects.create(
                    examination=examination,
                    method=form.cleaned_data['method'],
                    amount=form.cleaned_data['amount'],
                    status=form.cleaned_data['status']
                )

                # Generate the document
                generate_examination_document(examination)

                return redirect('employee_examination')  # Redirect to success page
            except Exception as e:
                return render(request, 'employee/add_examination.html', {
                    'form': form,
                    'account': account,
                    'error': f"An error occurred: {e}"
                })
    else:
        form = ExaminationForm()

    return render(request, 'employee/add_examination.html', {
        'form': form,
        'account': account,
    })

def generate_examination_document(examination):
    template_path = 'templates/examination_template.docx'
    output_path = os.path.join(settings.MEDIA_ROOT, 'examination_documents')  # Store the relative path

    if not os.path.exists(output_path):
        os.makedirs(output_path)

    # Load the document template
    doc = Document(template_path)

    # Populate the template with examination details
    patient = examination.patient
    doctor = examination.attending_doctor
    patient_full_name = patient.get_full_name_with_middle_initial()
    doctor_full_name = doctor.get_full_name_with_middle_initial()
    file_number = examination.get_file_number()
    patient_full_name_lastname_starting = patient.patient_full_name_last_name_start()

    # Generate the unique code
    salt = "CHMC2024"
    pepper = "WBMS2025"
    raw_code = f"{file_number}-{patient_full_name}-{doctor_full_name}"
    sha_input = f"{salt}{raw_code}{pepper}".encode('utf-8')
    unique_code = hashlib.sha256(sha_input).hexdigest()[:8].upper()
    unique_code = f"CHMC-{unique_code}"

    # Fill the template with text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace('{PATIENT_NAME}', patient_full_name)
                cell.text = cell.text.replace('{AGE}', str(patient.age))
                cell.text = cell.text.replace('{SEX}', patient.sex)
                cell.text = cell.text.replace('{SERVICE_TYPE}', ', '.join([str(s) for s in examination.service_types.all()]))
                cell.text = cell.text.replace('{DATE}', examination.date_created.strftime('%B %d, %Y'))
                cell.text = cell.text.replace('{FILE_NO}', file_number)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                # Add patient image if placeholder is found
                    if '{PATIENT_IMAGE}' in cell.text:
                        cell.text = cell.text.replace('{PATIENT_IMAGE}', '')  # Clear the placeholder text
                        if patient.image:
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(patient.image.path, width=Inches(1), height=Inches(1))

    footer = doc.sections[0].footer  # Get the footer from the first section

    # Iterate over paragraphs in the footer
    for paragraph in footer.paragraphs:
        # Replace placeholders in the footer
        if '{DOCTOR_NAME}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{DOCTOR_NAME}', doctor_full_name)

        if '{SIGNATURE}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{SIGNATURE}', '')  # Clear the placeholder
            if doctor.signature_image:
                run = paragraph.add_run()
                run.add_picture(doctor.signature_image.path, width=Inches(1), height=Inches(1))  # Add signature image

        if '{UNIQUE_CODE}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{UNIQUE_CODE}', unique_code)

    # Save the document
    output_filename = f"{patient_full_name_lastname_starting}.docx"
    output_full_path = os.path.join(output_path, output_filename)
    doc.save(output_full_path)

    # Attach the document to the Examination instance with the relative path
    examination.document.name = os.path.relpath(output_full_path, settings.MEDIA_ROOT).replace(os.sep, '/')
    examination.save()


@login_required
def upload_edited_document(request, pk):
    examination = get_object_or_404(Examination, pk=pk)

    if request.method == 'POST':
        form = UploadEditedDocumentForm(request.POST, request.FILES, instance=examination)
        if form.is_valid():
            # Delete the old document if it exists
            if examination.edited_document:
                if default_storage.exists(examination.edited_document.name):
                    default_storage.delete(examination.edited_document.name)

            # Save the new document
            form.save()

            user = request.user
            if user.is_superuser:
                return redirect('admin_examination')
            elif hasattr(user, 'is_associated_doctor') and user.is_associated_doctor:
                return redirect('assocdoc_examination')
            else:
                return redirect('employee_examination')
        else:
            return JsonResponse({'message': 'Failed to upload document.'}, status=400)

    return JsonResponse({'message': 'Invalid request method.'}, status=400)

def docx_to_pdf_exact(docx_path):

    try:
        # Initialize COM library
        pythoncom.CoInitialize()

        # Start MS Word
        word = Dispatch('Word.Application')
        word.Visible = False

        doc = word.Documents.Open(docx_path)
        pdf_path = docx_path.replace('.docx', '.pdf')
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 corresponds to the PDF format
        doc.Close()
        word.Quit()

        # Return the path to the PDF
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"Error during conversion: {e}")
    finally:
        # Ensure COM library is uninitialized
        pythoncom.CoUninitialize()

@login_required
def view_document(request, pk):

    examination = get_object_or_404(Examination, pk=pk)

    # Ensure an edited document exists
    if not examination.edited_document:
        user = request.user
        if user.is_superuser:
            return redirect('admin_examination')
        elif hasattr(user, 'is_associated_doctor') and user.is_associated_doctor:
            return redirect('assocdoc_examination')
        else:
            return redirect('employee_examination')

    # Convert .docx to PDF
    try:
        pdf_path = docx_to_pdf_exact(examination.edited_document.path)

        # Return the PDF file
        return FileResponse(open(pdf_path, 'rb'), content_type='application/pdf')
    except Exception as e:
        # Handle errors gracefully
        return HttpResponse(f"Error during document view: {e}", status=500)
    
def docx_to_pdf_auth_checker(docx_path):
    """
    Converts a .docx file to PDF using MS Word (COM automation).
    """
    try:
        pythoncom.CoInitialize()  # Initialize COM library
        word = Dispatch('Word.Application')
        word.Visible = False  # Keep Word invisible
        doc = word.Documents.Open(docx_path)
        pdf_path = docx_path.replace('.docx', '.pdf')
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 corresponds to PDF format
        doc.Close()
        word.Quit()
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"Error during conversion: {e}")
    finally:
        pythoncom.CoUninitialize()  # Cleanup COM library

def verify_document(request):
    """
    Verify a document by checking the unique code and displaying the PDF if valid.
    """
    if request.method == 'POST':
        unique_code = request.POST.get('unique_code')

        try:
            # Split the code to get necessary parts
            unique_code_parts = unique_code.split('-')
            if len(unique_code_parts) != 2:
                return HttpResponse("Invalid code format.")

            code_prefix, code_value = unique_code_parts
            if code_prefix != "CHMC":
                return HttpResponse("Invalid prefix in the code.")

            # Iterate through all examinations to find the one that matches the unique code
            examination = None
            for ex in Examination.objects.all():
                if ex.get_unique_code() == unique_code:
                    examination = ex
                    break

            if not examination:
                return HttpResponse("Document not found for the provided code.")

            # Check if an edited document exists
            if not examination.edited_document:
                return HttpResponse("This document does not have an edited version.")
            if not os.path.exists(examination.edited_document.path):
                return HttpResponse("Edited document file is missing.")

            # Convert the .docx document to PDF if not already converted
            try:
                pdf_path = docx_to_pdf_auth_checker(examination.edited_document.path)

                # Serve the PDF file for the user to view
                return FileResponse(open(pdf_path, 'rb'), content_type='application/pdf')
            except Exception as e:
                return HttpResponse(f"Error during document conversion: {e}")
        
        except Exception as e:
            return HttpResponse(f"Error: {e}")
    else:
        return render(request, 'authenticity_checker.html')

    
@csrf_exempt
def search_patient(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            query = data.get("query", "").strip()
            if query:
                # Split the query into parts (e.g., "John Doe")
                name_parts = query.split()
                filters = Q()

                if len(name_parts) == 1:
                    # Search for matches in first, middle, or last name
                    filters |= (
                        Q(first_name__icontains=name_parts[0]) |
                        Q(middle_name__icontains=name_parts[0]) |
                        Q(last_name__icontains=name_parts[0])
                    )
                elif len(name_parts) > 1:
                    # Search for first + last name or other combinations
                    filters |= (
                        Q(first_name__icontains=name_parts[0]) & 
                        Q(last_name__icontains=name_parts[1])
                    )
                    # Optionally include middle name if three parts are provided
                    if len(name_parts) > 2:
                        filters |= (
                            Q(first_name__icontains=name_parts[0]) &
                            Q(middle_name__icontains=name_parts[1]) &
                            Q(last_name__icontains=name_parts[2])
                        )

                # Filter patients based on the query
                patients = Patient.objects.filter(filters)[:10]
                results = [
                    {
                        "id": patient.id,
                        "first_name": patient.first_name or "",
                        "middle_name": patient.middle_name or "",
                        "last_name": patient.last_name or "",
                        "age" : patient.age or "",
                        "sex" : patient.sex or "",
                        "contact_number" : patient.contact_number or "",
                        "address" :patient.address or ""
                    }
                    for patient in patients
                ]
                return JsonResponse({"patients": results})

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format"}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"patients": []})

@login_required
def upload_examination_result_image(request, pk):
    if request.method == "POST":
        exam = get_object_or_404(Examination, pk=pk)
        form = UploadResultImageForm(request.POST, instance=exam)
        if form.is_valid():
            image_result = request.POST.get('result_image')
            if image_result:
                        # Convert base64 image data to a Django ImageFile
                format, imgstr = image_result.split(';base64,')  # Get base64 string
                ext = format.split('/')[1]  # Extract file extension (png, jpeg, etc.)
                image_result = ContentFile(base64.b64decode(imgstr), name=f"examination_{exam.id}_image.{ext}")
                        
                        # Save the new image to the patient instance
                exam.result_image = image_result
                exam.save()
                user = request.user
                if user.is_superuser:
                    return redirect('admin_examination')
                elif hasattr(user, 'is_associated_doctor') and user.is_associated_doctor:
                    return redirect('assocdoc_examination')
                else:
                    return redirect('employee_examination')
        else:
            return JsonResponse({"status": "error", "message": "Invalid data."})
    user = request.user
    if user.is_superuser:
        return redirect('admin_examination')
    elif hasattr(user, 'is_associated_doctor') and user.is_associated_doctor:
        return redirect('assocdoc_examination')
    else:
        return redirect('employee_examination')

@login_required
def edit_examination(request, pk):
    account = request.user
    exam = get_object_or_404(Examination, pk=pk)
    payment = exam.payment_set.first()  # Assuming one payment per exam

    PAYMENT_METHODS = Payment.PAYMENT_METHODS
    if request.method == 'POST':
        form = EditExaminationForm(request.POST)
        if form.is_valid():
            # Update patient details
            if form.cleaned_data.get('patient_first_name'):
                exam.patient.first_name = form.cleaned_data['patient_first_name']
            if form.cleaned_data.get('patient_middle_name'):
                exam.patient.middle_name = form.cleaned_data['patient_middle_name']
            if form.cleaned_data.get('patient_last_name'):
                exam.patient.last_name = form.cleaned_data['patient_last_name']
            exam.patient.save()

            # Update service types
            if form.cleaned_data.get('service_types'):
                exam.service_types.set(form.cleaned_data['service_types'])

            # Update payment details
            if payment:
                payment.method = form.cleaned_data.get('payment_method', payment.method)
                payment.status = form.cleaned_data.get('payment_status', payment.status)
                payment.amount = form.cleaned_data.get('payment_amount', payment.amount)
                payment.save()

            # Use 'next' GET parameter to determine whether to redirect or render
            next_template = request.GET.get('next_template')
            if not next_template:
                if account.is_superuser:
                    next_template = 'admin/examination.html'
                elif hasattr(account, 'is_associated_doctor') and account.is_associated_doctor:
                    next_template = 'associated_doctor/examination.html'
                else:
                    next_template = 'employee/examination.html'
            
            # Check if the next_template is a valid template or return redirect if necessary
            if next_template.endswith('.html'):
                examinations = Examination.objects.all().order_by('-date_created')
                return render(request, next_template, {
                    'form': form, 
                    'exam': exam,
                    'account' : account,  
                    'PAYMENT_METHODS' : PAYMENT_METHODS,
                    'examinations' : examinations,
                    })
            else:
                return redirect(next_template)
        else:
            return JsonResponse({'success': False, 'errors': form.errors})

    # Prepopulate the form for GET requests
    initial_data = {
        'patient_first_name': exam.patient.first_name,
        'patient_middle_name': exam.patient.middle_name,
        'patient_last_name': exam.patient.last_name,
        'service_types': exam.service_types.all(),
        'payment_method': payment.method if payment else '',
        'payment_status': payment.status if payment else '',
        'payment_amount': payment.amount if payment else '',
        
    }
    form = EditExaminationForm(initial=initial_data)

    # Use the next_template GET parameter to determine which template to render
    next_template = request.GET.get('next_template')
    if not next_template:
        if account.is_superuser:
            next_template = 'admin/admin_examination.html'
        elif hasattr(account, 'is_associated_doctor') and account.is_associated_doctor:
            next_template = 'associated_doctor/examination.html'
        else:
            next_template = 'employee/examination.html'

    return render(request, next_template, {'form': form, 'exam': exam, 'account': account, 'PAYMENT_METHODS' : PAYMENT_METHODS})

@login_required
def edit_patient(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)
    if request.method == 'POST':
        form = PatientEditForm(request.POST, request.FILES, instance=patient)
        if form.is_valid():
            form.save()
            return redirect('employee_patients_list_with_id', patient_id=patient.id)  # Redirect after saving changes
    else:
        form = PatientEditForm(instance=patient)
    
    account = request.user
    context = {
        'form': form,
        'patient': patient,
        'account': account,
    }
    return render(request, 'employee/edit_patient.html', context)

@user_passes_test(lambda u: u.is_associated_doctor)
def assocdoc_dashboard_view(request):
    account = request.user
    return render(request, 'associated_doctor/assocdoc_dashboard.html', {'account': account})